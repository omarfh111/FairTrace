"""
Generate FairTrace Technical Report as Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading(doc, text, level):
    """Add a heading with custom styling."""
    heading = doc.add_heading(text, level)
    return heading

def add_code_block(doc, code):
    """Add a code block with monospace font."""
    para = doc.add_paragraph()
    run = para.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.5)
    return para

def create_technical_report():
    doc = Document()
    
    # Title
    title = doc.add_heading('FairTrace - Multi-Agent Credit Decision System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Technical Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph('')  # Spacer
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(
        'FairTrace is a production-grade, multi-agent AI system designed for explainable credit decisioning. '
        'The system evaluates credit applications (consumer loans, startup funding, enterprise credit) using a '
        'debate-based architecture where multiple specialized agents analyze each application from different '
        'perspectives. An orchestrator then synthesizes their verdicts into a final, explainable decision.'
    )
    
    # Key Features table
    doc.add_heading('Key Features', 2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    features = [
        ('Multi-Agent Debate', 'Risk, Fairness, and Trajectory agents debate each application'),
        ('Parallel Execution', 'Agents run concurrently using async LangGraph (~10s total)'),
        ('Hybrid Vector Search', 'Dense + sparse embeddings with Qdrant for case retrieval'),
        ('Persistent Storage', 'Supabase PostgreSQL for decisions and agent cache'),
        ('On-Demand Agents', 'Advisor, Narrative, Comparator, Scenario agents for deeper insights'),
        ('Full Observability', 'LangSmith tracing for every LLM call and retrieval'),
        ('Modern React UI', 'Real-time dashboard with evidence visualization'),
    ]
    for i, (feature, desc) in enumerate(features):
        table.rows[i].cells[0].text = feature
        table.rows[i].cells[1].text = desc
    
    doc.add_page_break()
    
    # System Architecture
    doc.add_heading('System Architecture', 1)
    
    doc.add_heading('High-Level Overview', 2)
    doc.add_paragraph(
        'The FairTrace system consists of three main layers: a React frontend for user interaction, '
        'a FastAPI backend that orchestrates the multi-agent workflow, and external services for '
        'data persistence and observability.'
    )
    
    # Architecture Components
    doc.add_heading('Frontend Layer', 3)
    doc.add_paragraph('• Dashboard View - Main decision interface with application form')
    doc.add_paragraph('• Evidence Panel - Displays similar cases retrieved by agents')
    doc.add_paragraph('• Decision Details - Full breakdown of agent verdicts')
    doc.add_paragraph('• On-Demand Agent Panels - Advisor, Narrative, Comparator, Scenario agents')
    
    doc.add_heading('Backend Layer (FastAPI)', 3)
    doc.add_paragraph('• REST API endpoints for decisions and on-demand agents')
    doc.add_paragraph('• LangGraph workflow for parallel agent orchestration')
    doc.add_paragraph('• Async execution using asyncio and ainvoke()')
    doc.add_paragraph('• Centralized configuration via Pydantic settings')
    
    doc.add_heading('External Services', 3)
    doc.add_paragraph('• Supabase PostgreSQL - Persistent storage for decisions and agent cache')
    doc.add_paragraph('• Qdrant Cloud - Vector database with 1000+ synthetic cases')
    doc.add_paragraph('• LangSmith - Full observability and tracing')
    doc.add_paragraph('• OpenAI - GPT-4o-mini for agent reasoning')
    doc.add_paragraph('• Ollama - Local embeddings (mxbai-embed-large)')
    
    doc.add_page_break()
    
    # RAG Pipeline
    doc.add_heading('RAG Pipeline Architecture', 1)
    
    doc.add_paragraph(
        'Each agent uses a sophisticated Retrieval-Augmented Generation (RAG) pipeline that combines '
        'dense and sparse embeddings for optimal retrieval quality.'
    )
    
    doc.add_heading('Step 1: Query Generation', 2)
    doc.add_paragraph(
        'The application data is processed by GPT-4o-mini to generate a semantic search query '
        'that captures the essence of the application for similarity matching.'
    )
    
    doc.add_heading('Step 2: Hybrid Embedding', 2)
    doc.add_paragraph('The query is embedded using two complementary approaches:')
    doc.add_paragraph('')
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Embedding Type'
    table.rows[0].cells[1].text = 'Model'
    table.rows[0].cells[2].text = 'Purpose'
    table.rows[1].cells[0].text = 'Dense'
    table.rows[1].cells[1].text = 'Ollama mxbai-embed-large (1024 dim)'
    table.rows[1].cells[2].text = 'Semantic meaning ("financial stability" ≈ "economic security")'
    table.rows[2].cells[0].text = 'Sparse'
    table.rows[2].cells[1].text = 'FastEmbed BM25'
    table.rows[2].cells[2].text = 'Exact terms, acronyms, domain jargon ("DTI" = "DTI")'
    
    doc.add_heading('Step 3: RRF Fusion', 2)
    doc.add_paragraph(
        'Reciprocal Rank Fusion (RRF) combines dense and sparse scores to leverage the strengths '
        'of both approaches. This ensures both semantic similarity and keyword matching contribute '
        'to the final ranking.'
    )
    
    doc.add_heading('Step 4: Qdrant Vector Search', 2)
    doc.add_paragraph('The fairtrace_cases collection contains 1000+ synthetic cases across three types:')
    doc.add_paragraph('• Client Cases (500+): Consumer loans, mortgages, personal credit')
    doc.add_paragraph('• Startup Cases (300+): VC funding rounds, seed/Series A/B')
    doc.add_paragraph('• Enterprise Cases (200+): Credit lines, working capital, trade finance')
    doc.add_paragraph('')
    doc.add_paragraph('Each case contains:')
    doc.add_paragraph('• Application data (metrics, financials)')
    doc.add_paragraph('• Historical outcome (APPROVED, REJECTED, DEFAULT, PERFORMING, etc.)')
    doc.add_paragraph('• Dense + Sparse embedding vectors')
    
    doc.add_heading('Step 5: Context Augmentation', 2)
    doc.add_paragraph(
        'Retrieved cases (top 5-10) are formatted into a context block that provides the LLM '
        'with relevant historical precedents for making evidence-based decisions.'
    )
    
    doc.add_heading('Step 6: LLM Generation', 2)
    doc.add_paragraph('OpenAI GPT-4o-mini generates the agent verdict with:')
    doc.add_paragraph('• Structured JSON output (json_object mode)')
    doc.add_paragraph('• Temperature: 0 for consistency')
    doc.add_paragraph('• Full tracing via LangSmith')
    
    doc.add_page_break()
    
    # Agent Details
    doc.add_heading('Agent Architecture', 1)
    
    doc.add_heading('Core Decision Agents (Parallel Execution)', 2)
    doc.add_paragraph(
        'These three agents run simultaneously when a decision is requested. '
        'LangGraph orchestrates their parallel execution using async/await patterns.'
    )
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    headers = ['Agent', 'Role', 'Search Strategy', 'Key Metrics']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    
    agents = [
        ('Risk Agent', "Devil's advocate - finds reasons to reject", 
         'Search for similar cases that defaulted or were problematic', 
         'Risk level, red flags, mitigating factors'),
        ('Fairness Agent', 'Ensures consistency with similar approved cases', 
         'Search for approved cases with similar profiles', 
         'Consistency score, similar approved count'),
        ('Trajectory Agent', 'Predicts future outcomes based on patterns', 
         'Search for cases with similar starting points', 
         'Growth pattern, predicted outcome probability'),
    ]
    for i, (name, role, strategy, metrics) in enumerate(agents, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = role
        table.rows[i].cells[2].text = strategy
        table.rows[i].cells[3].text = metrics
    
    doc.add_heading('Orchestrator', 2)
    doc.add_paragraph(
        'After all three agents complete, the Orchestrator synthesizes their verdicts into a '
        'final decision. It weighs each agent\'s recommendation, confidence, and supporting evidence '
        'to produce one of four outcomes: APPROVE, CONDITIONAL, REJECT, or ESCALATE.'
    )
    
    doc.add_heading('On-Demand Agents (Lazy-Loaded)', 2)
    doc.add_paragraph(
        'These agents are invoked only when the user requests additional analysis. '
        'Results are cached in the database for efficiency.'
    )
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    headers = ['Agent', 'Trigger', 'Search Strategy', 'Output']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    
    ondemand = [
        ('Advisor Agent', 'User clicks "Get Recommendations"', 
         'Search for improved cases that got approved after changes', 
         'Specific, actionable improvement steps'),
        ('Narrative Agent', 'User clicks "See Stories"', 
         'Search for notable success/failure stories', 
         'Compelling narratives with lessons learned'),
        ('Comparator Agent', 'User clicks "Gap Analysis"', 
         'Search for top approved cases in same category', 
         'Metric-by-metric comparison with benchmarks'),
        ('Scenario Agent', 'User defines what-if scenarios', 
         'Re-evaluate with modified application data', 
         'Probability changes, optimal path to approval'),
    ]
    for i, (name, trigger, strategy, output) in enumerate(ondemand, 1):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = trigger
        table.rows[i].cells[2].text = strategy
        table.rows[i].cells[3].text = output
    
    doc.add_page_break()
    
    # API Reference
    doc.add_heading('API Reference', 1)
    
    doc.add_paragraph('Base URL: http://localhost:8000/api/v1')
    doc.add_paragraph('')
    
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Method'
    table.rows[0].cells[1].text = 'Endpoint'
    table.rows[0].cells[2].text = 'Description'
    
    endpoints = [
        ('POST', '/decisions', 'Submit application for decision'),
        ('GET', '/decisions/{id}', 'Retrieve decision by ID'),
        ('GET', '/decisions', 'List recent decisions'),
        ('GET', '/decisions/{id}/advisor', 'Get improvement recommendations'),
        ('GET', '/decisions/{id}/narrative', 'Get historical narratives'),
        ('GET', '/decisions/{id}/comparator', 'Get gap analysis'),
        ('POST', '/decisions/{id}/scenario', 'Run what-if scenarios'),
        ('GET', '/health', 'Health check'),
    ]
    for i, (method, endpoint, desc) in enumerate(endpoints, 1):
        table.rows[i].cells[0].text = method
        table.rows[i].cells[1].text = endpoint
        table.rows[i].cells[2].text = desc
    
    doc.add_heading('Example Request', 2)
    add_code_block(doc, '''POST /api/v1/decisions
Content-Type: application/json

{
  "application": {
    "age": 35,
    "contract_type": "CDI",
    "income_annual": 55000,
    "debt_to_income_ratio": 0.35,
    "missed_payments_last_12m": 1,
    "loan_purpose": "Home improvement"
  }
}''')
    
    doc.add_heading('Example Response', 2)
    add_code_block(doc, '''{
  "decision_id": "abc123",
  "application_type": "client",
  "final_decision": {
    "recommendation": "CONDITIONAL",
    "confidence": "MEDIUM",
    "risk_level": "MEDIUM",
    "reasoning": "The application shows stable income but elevated DTI...",
    "conditions": ["Reduce debt-to-income below 0.30"],
    "key_factors": ["DTI ratio", "Payment history"]
  },
  "risk_verdict": {...},
  "fairness_verdict": {...},
  "trajectory_verdict": {...},
  "processing_time_ms": 10234
}''')
    
    doc.add_page_break()
    
    # Project Structure
    doc.add_heading('Project Structure', 1)
    
    add_code_block(doc, '''fairtrace/
├── api/                        # FastAPI backend
│   ├── main.py                 # Application entry point
│   ├── schemas.py              # Pydantic models
│   └── routes/
│       └── decisions.py        # API endpoints
│
├── agents/                     # AI agents
│   ├── base_agent.py           # Shared LLM configuration
│   ├── risk_agent.py           # Risk assessment
│   ├── fairness_agent.py       # Fairness evaluation
│   ├── trajectory_agent.py     # Outcome prediction
│   ├── advisor_agent.py        # Improvement recommendations
│   ├── narrative_agent.py      # Historical narratives
│   ├── comparator_agent.py     # Gap analysis
│   └── scenario_agent.py       # What-if scenarios
│
├── graph/                      # LangGraph workflow
│   └── decision_graph.py       # Parallel agent orchestration
│
├── tools/                      # Agent tools
│   └── qdrant_retriever.py     # Hybrid vector search
│
├── db/                         # Database layer
│   ├── __init__.py
│   └── repository.py           # Async CRUD operations
│
├── data_generation/            # Synthetic data generation
│   ├── generate_data.py        # Main data generator script
│   ├── prompts_config.py       # LLM prompts for data generation
│   └── output/                 # Generated CSV/JSON files
│
├── ingestion/                  # Vector DB ingestion
│   └── ingest_to_qdrant.py     # Upload data to Qdrant
│
├── evaluation/                 # Evaluation framework
│   ├── generate_eval_dataset.py
│   ├── run_evaluation.py
│   └── metrics/
│
├── frontend/                   # React frontend
│   ├── src/
│   │   ├── components/         # UI components
│   │   ├── lib/api.ts          # API client
│   │   └── types/              # TypeScript types
│   └── package.json
│
├── config.py                   # Centralized configuration
├── requirements.txt            # Python dependencies
└── .env.example                # Environment template''')
    
    doc.add_page_break()
    
    # Technology Stack
    doc.add_heading('Technology Stack', 1)
    
    doc.add_heading('Backend', 2)
    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Component'
    table.rows[0].cells[1].text = 'Technology'
    table.rows[0].cells[2].text = 'Version'
    
    tech = [
        ('Web Framework', 'FastAPI', '0.115+'),
        ('Agent Orchestration', 'LangGraph', '0.2+'),
        ('LLM Framework', 'LangChain', '0.3+'),
        ('LLM', 'OpenAI GPT-4o-mini', 'Latest'),
        ('Embeddings', 'Ollama mxbai-embed-large', '1024 dim'),
        ('Vector Database', 'Qdrant Cloud', 'Latest'),
        ('SQL Database', 'Supabase PostgreSQL', '15+'),
    ]
    for i, (comp, tech_name, ver) in enumerate(tech, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = tech_name
        table.rows[i].cells[2].text = ver
    
    doc.add_heading('Frontend', 2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Component'
    table.rows[0].cells[1].text = 'Technology'
    table.rows[0].cells[2].text = 'Version'
    
    fe_tech = [
        ('Framework', 'React', '18+'),
        ('Build Tool', 'Vite', '5+'),
        ('Language', 'TypeScript', '5+'),
        ('UI Components', 'shadcn/ui + Radix', 'Latest'),
    ]
    for i, (comp, tech_name, ver) in enumerate(fe_tech, 1):
        table.rows[i].cells[0].text = comp
        table.rows[i].cells[1].text = tech_name
        table.rows[i].cells[2].text = ver
    
    doc.add_page_break()
    
    # Performance
    doc.add_heading('Performance Characteristics', 1)
    
    doc.add_heading('Decision Latency', 2)
    doc.add_paragraph('• Total decision time: ~10 seconds (with parallel agent execution)')
    doc.add_paragraph('• Agent execution: ~4-5 seconds each (running in parallel)')
    doc.add_paragraph('• Orchestrator synthesis: ~2 seconds')
    doc.add_paragraph('• Vector search: <500ms per agent')
    
    doc.add_heading('Scalability', 2)
    doc.add_paragraph('• Async execution: Agents run concurrently using asyncio.to_thread()')
    doc.add_paragraph('• Database: PostgreSQL with async driver (asyncpg)')
    doc.add_paragraph('• Caching: On-demand agent results cached in database')
    doc.add_paragraph('• Vector DB: Qdrant Cloud with horizontal scaling support')
    
    doc.add_heading('Observability', 2)
    doc.add_paragraph('LangSmith provides full observability:')
    doc.add_paragraph('• LLM call traces with prompts and responses')
    doc.add_paragraph('• Token usage and cost tracking')
    doc.add_paragraph('• Latency breakdown by component')
    doc.add_paragraph('• Error tracking and debugging')
    
    doc.add_page_break()
    
    # Evaluation
    doc.add_heading('Evaluation Framework', 1)
    
    doc.add_paragraph('The evaluation framework measures three key dimensions:')
    
    doc.add_heading('1. Retrieval Quality', 2)
    doc.add_paragraph('• Are the right cases being retrieved for each query?')
    doc.add_paragraph('• Measured via precision@k and recall@k on golden dataset')
    
    doc.add_heading('2. Decision Consistency', 2)
    doc.add_paragraph('• Do similar applications get similar decisions?')
    doc.add_paragraph('• Measured via consistency score across synthetic test cases')
    
    doc.add_heading('3. Explanation Quality', 2)
    doc.add_paragraph('• Are the explanations coherent and helpful?')
    doc.add_paragraph('• Evaluated using LLM-as-judge with structured criteria')
    
    doc.add_heading('Running Evaluation', 2)
    add_code_block(doc, '''# Generate evaluation dataset
python evaluation/generate_eval_dataset.py

# Run evaluation
python evaluation/run_evaluation.py''')
    
    # Conclusion
    doc.add_page_break()
    doc.add_heading('Conclusion', 1)
    doc.add_paragraph(
        'FairTrace represents a production-ready implementation of explainable AI for credit decisioning. '
        'The multi-agent debate architecture ensures decisions are well-reasoned and supported by evidence '
        'from similar historical cases. Key innovations include:'
    )
    doc.add_paragraph('• Parallel agent execution for reduced latency')
    doc.add_paragraph('• Hybrid dense+sparse embeddings for optimal retrieval')
    doc.add_paragraph('• Full persistence with Supabase PostgreSQL')
    doc.add_paragraph('• Comprehensive observability via LangSmith')
    doc.add_paragraph('• On-demand agents for deeper analysis without upfront cost')
    
    # Save
    doc.save('docs/FairTrace_Technical_Report.docx')
    print('✅ Technical report saved to docs/FairTrace_Technical_Report.docx')

if __name__ == '__main__':
    import os
    os.makedirs('docs', exist_ok=True)
    create_technical_report()
